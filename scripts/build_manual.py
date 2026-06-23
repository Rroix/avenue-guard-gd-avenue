from __future__ import annotations

import ast
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCX_PATH = DOCS / "Avenue_Guard_Manual.docx"
MD_PATH = DOCS / "Avenue_Guard_Manual.md"
PDF_PATH = DOCS / "Avenue_Guard_Manual.pdf"

BLUE = RGBColor(0x1D, 0x6B, 0x73)
DARK_BLUE = RGBColor(0x17, 0x3F, 0x46)
INK = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
TABLE_FILL = "DCEDEA"
CALLOUT_FILL = "FFF7E6"
CODE_FILL = "F6F7F8"


def p(text: str) -> tuple:
    return ("p", text)


def h2(text: str) -> tuple:
    return ("h2", text)


def h3(text: str) -> tuple:
    return ("h3", text)


def bullets(items: list[str]) -> tuple:
    return ("bullets", items)


def numbered(items: list[str]) -> tuple:
    return ("numbered", items)


def callout(title: str, text: str) -> tuple:
    return ("callout", title, text)


def table(headers: list[str], rows: list[list[str]]) -> tuple:
    return ("table", headers, rows)


def diagram(title: str, steps: list[str], note: str = "") -> tuple:
    return ("diagram", title, steps, note)


def code(title: str, language: str, text: str) -> tuple:
    return ("code", title, language, text.rstrip())


def source_code(title: str, rel_path: str, start: int, end: int, language: str = "python") -> tuple:
    path = ROOT / rel_path
    try:
        lines = path.read_text(encoding="utf-8").splitlines()
        selected = lines[max(0, start - 1): max(start - 1, end)]
        numbered = [f"{idx:>4}: {line}" for idx, line in enumerate(selected, start=start)]
        text = "\n".join(numbered)
    except Exception as exc:
        text = f"# Could not load {rel_path}:{start}-{end}: {type(exc).__name__}"
    return code(f"{title} ({rel_path}:{start}-{end})", language, text)


CHAPTERS = [
    (
        "How To Read This Manual",
        [
            p(
                "Avenue Guard is the operating system for GD Avenue's Discord workflows. It is not only a moderation bot, "
                "and it is not only a level request bot. It connects request waves, weekly activity rewards, tickets, "
                "help flows, forum formatting, staff logs, analytics, and admin diagnostics into one persistent bot."
            ),
            p(
                "This manual is private technical documentation for Rodrigo. It is not written as a staff handbook. It explains "
                "how the bot is built, how the code thinks, which modules own which workflows, and why the implementation choices "
                "exist. Staff workflows are described only because they are part of the bot's code and state model."
            ),
            callout(
                "Core idea",
                "The bot is built around a single configured Discord server, a JSON configuration file, persistent SQLite "
                "state, and a set of cogs that each own a clear part of the community workflow.",
            ),
            h2("Document Map"),
            bullets(
                [
                    "The first chapters give the big mental model: runtime, config, persistence, cogs, permissions, and history.",
                    "The workflow chapters explain requests, weekly tracking, help/tickets, forums, telemetry, admin tools, and impact reporting.",
                    "The private deep-dive chapters walk through actual source code excerpts and the less obvious engineering ideas.",
                    "The appendices give quick lookup tables for commands, database tables, files, and maintenance habits.",
                ]
            ),
            p(
                "The bot has grown through many rounds of practical server needs. That matters because many design choices "
                "are not abstract engineering preferences. They exist because the community needed safer requesting, better "
                "review visibility, more reliable tickets, measurable activity, and staff tools that can recover from missing "
                "messages or restarts."
            ),
        ],
    ),
    (
        "Executive Overview",
        [
            p(
                "At the highest level, Avenue Guard is a Discord operations bot. It listens to server events, direct messages, "
                "button interactions, modal submissions, slash commands, scheduled loops, and background telemetry. Each event "
                "is routed through a cog that owns the relevant workflow, and most important outcomes are stored in SQLite."
            ),
            p(
                "The bot's central value is continuity. A request wave should not disappear after a restart. A ticket should "
                "have an ID and a transcript. A weekly reward should know who was contacted and whether they replied. A forum "
                "post deleted for missing the required word should be logged. A reviewer should see which requests are still "
                "pending. An owner should be able to generate numbers that show the bot's impact."
            ),
            diagram(
                "Main Operating Loop",
                [
                    "Discord event or command",
                    "Persistent view or cog",
                    "Config plus database state",
                    "Discord response",
                    "Logs, summaries, and metrics",
                ],
                "Every major workflow follows this loop so visible Discord state and stored bot memory stay aligned.",
            ),
            table(
                ["Area", "What Avenue Guard Does", "Why It Matters"],
                [
                    ["Request waves", "Opens, limits, schedules, validates, reviews, summarizes, and repairs level request waves.", "Turns a messy manual process into a controlled staff queue."],
                    ["Weekly rewards", "Tracks eligible activity, contacts winners, records claims, and routes weekly submissions into the review workflow.", "Rewards activity while keeping staff review consistent."],
                    ["Help and tickets", "Runs a DM help dashboard, FAQ search, appeals, reports, bot issues, transcript requests, tickets, and satisfaction prompts.", "Gives members private support without losing staff accountability."],
                    ["Guardrails", "Deletes restricted proof-channel misuse, applies restriction roles, sends role DMs, manages sticky/forum reminders, and enforces forum words.", "Reduces repeated moderation work and keeps public channels organized."],
                    ["Telemetry", "Tracks daily summaries, command usage, voice time, activity, anti-farm events, and now persistent impact exports.", "Makes the bot measurable and useful for operations, not just automation."],
                ],
            ),
            h2("What Makes It Complex"),
            p(
                "Avenue Guard is complex because it combines multiple stateful systems. It has persistent Discord views, "
                "scheduled jobs, slash commands, mod-only and admin-only gates, external validation calls, modal workflows, "
                "message edits, channel creation and deletion, ticket transcripts, and configurable embed templates. Many bots "
                "do one of these things. Avenue Guard coordinates all of them in one server-specific package."
            ),
            p(
                "The most important invariant is that public Discord state and database state should describe the same reality. "
                "If requests are closed in SQLite, the request button should say closed. If a ticket is resolved, the opening "
                "message and transcript should reflect that. If a request is reviewed, the original buttons should be disabled. "
                "Much of the bot's repair and diagnostic design exists to preserve this agreement between what users see and "
                "what the bot remembers."
            ),
        ],
    ),
    (
        "Development History",
        [
            p(
                "The current bot evolved from a simpler Render-ready Discord bot into a much more complete community operations "
                "platform. The early shape focused on practical automations: keeping the bot in the correct guild, deleting "
                "misplaced proof-channel messages, sending role-triggered DMs, and keeping sticky reminder messages visible."
            ),
            p(
                "The next major phase added weekly activity tracking. The bot began counting eligible member messages, skipping "
                "excluded roles and channels, and contacting weekly winners with request opportunities. This phase introduced "
                "the idea that activity and reward state must survive restarts, because weekly workflows can span hours or days."
            ),
            p(
                "The request system then became the largest feature area. Live request waves gained an open/closed state, "
                "count limits, timers, scheduled openings, request types, duplicate blocking, per-wave summaries, reviewer "
                "buttons, result channels, edit windows, request edit audits, validation cache, and repair commands. Weekly "
                "request submissions were later brought into the same review workflow so staff did not need to learn two systems."
            ),
            p(
                "The help system grew in parallel. Instead of only sending generic DMs, Avenue Guard now runs an in-DMs dashboard, "
                "FAQ search, pre-ticket FAQ suggestions, appeal/report/bug previews, staff reply relay, private ticket creation, "
                "ticket statuses, transcript saving, transcript search, transcript requests, and satisfaction prompts."
            ),
            p(
                "More recent phases focused on staff experience and measurement: cleaner log embeds, daily and weekly summaries, "
                "anti-farm detection, server icon rotation, admin dashboards, doctor/repair suggestions, and now impact reporting. "
                "The bot's direction has been consistent: when a process starts to need staff memory, the bot stores it, exposes it, "
                "and makes it reviewable."
            ),
            callout(
                "History in one sentence",
                "Avenue Guard developed from a utility bot into a persistent community workflow engine for requests, support, "
                "moderation guardrails, analytics, and staff coordination.",
            ),
        ],
    ),
    (
        "Runtime And Startup Architecture",
        [
            p(
                "The runtime begins in main.py. The bot creates a py-cord Bot object with the intents needed for messages, "
                "members, reactions, moderation, presences, voice states, and direct messages. It loads config.json through "
                "utils.config.Config, resolves the configured database path, opens it through utils.db.Database, installs "
                "global error handlers, and loads each cog extension."
            ),
            p(
                "Startup is deliberately defensive. On ready, the bot connects the database, checks that it can see the allowed "
                "guild, starts the keepalive server, starts background tasks in the tracking, help, request, and background cogs, "
                "and registers persistent views. Persistent views are crucial because Discord button interactions can arrive "
                "after a restart. The custom IDs live in utils.views and route back to the correct cog."
            ),
            table(
                ["Startup Step", "Owner", "Purpose"],
                [
                    ["Load config", "utils.config.Config", "Reads config.json and exposes typed getters for IDs, lists, and strings."],
                    ["Connect DB", "utils.db.Database", "Creates or migrates SQLite tables before workflows depend on them."],
                    ["Load cogs", "main.py", "Attaches feature modules for moderation, tracking, help, responses, sticky messages, requests, commands, and background jobs."],
                    ["Start tasks", "on_ready", "Starts loops for weekly handling, ticket scans, request auto-close, scheduled openings, summaries, status, and icon rotation."],
                    ["Register views", "utils.views", "Keeps buttons and selects alive across restarts through stable custom IDs."],
                ],
            ),
            h2("Hosted Environment"),
            p(
                "The bot is designed to run on a hosted service such as Render. A small keepalive HTTP server exists for hosted "
                "environments, but the important persistence requirement is the SQLite database. If the host uses ephemeral "
                "storage, the database path must live on a persistent disk or equivalent mounted storage."
            ),
            p(
                "The bot also assumes that startup may happen after Discord components already exist. That is why persistent "
                "views are registered every time the bot becomes ready and why request/ticket state is not reconstructed from "
                "memory. Startup should be safe whether the bot was restarted manually, redeployed by the host, or recovered "
                "after an exception."
            ),
        ],
    ),
    (
        "Configuration Model",
        [
            p(
                "Avenue Guard uses config.json as the main control plane. The Config loader treats keys beginning with an "
                "underscore as comments by convention, but it does not need a separate schema file. IDs can be stored as strings; "
                "the getter methods convert them to integers or lists when code needs them."
            ),
            p(
                "This design makes server-specific changes practical. Admins can change channel IDs, role IDs, request wording, "
                "embed templates, validation settings, icon URLs, background summary settings, and help FAQ entries without "
                "editing Python code. The /resync command reloads config and response rules so many changes do not require a full restart."
            ),
            table(
                ["Config Section", "Controls"],
                [
                    ["guild", "The allowed guild ID. The bot shuts down if it cannot operate in that server."],
                    ["roles", "Moderation, admin, tracking exclusion, reward, and watched-role IDs."],
                    ["channels", "Core log, request, transcript, command, proof, and help channels."],
                    ["tracking", "Weekly message counting, winner DMs, reminders, streaks, anti-farm checks, and logs."],
                    ["tickets", "Ticket category, staff ping role, cooldowns, inactivity, and satisfaction prompt."],
                    ["help", "FAQ entries, warnings, cooldown assumptions, duplicate windows, and submission limits."],
                    ["level_requests", "Request channels, roles, text, embeds, validation, wave summaries, colors, and opening announcements."],
                    ["background", "Daily summaries, weekly recaps, rotating status, and server icon rotation."],
                    ["database", "SQLite path and scheduled zipped database backups."],
                    ["impact", "Destination for persistent impact report attachments and snapshots."],
                ],
            ),
            h2("Why JSON Instead Of Hardcoding"),
            p(
                "The server changes faster than code should. Roles are renamed, channels are moved, request copy is adjusted, "
                "and staff may want different embed language for review or result messages. Keeping those values in config.json "
                "lets the bot remain stable while the server's surface changes."
            ),
            p(
                "The most customizable parts are the embed templates. Request submissions, reviewed requests, weekly submissions, "
                "result messages, wave summaries, help logs, and announcements are built from template variables. This gives the "
                "server control over tone and layout while keeping the workflow rules in code. The config checker validates many "
                "of those templates so a typo in a variable is easier to catch before staff depend on the embed."
            ),
        ],
    ),
    (
        "Database And Persistence",
        [
            p(
                "The SQLite layer is intentionally small and predictable. utils.db.Database owns one SQLite connection with "
                "check_same_thread disabled, serializes operations with an asyncio lock, and runs blocking database work inside "
                "threads. The migration code creates tables and adds columns for older databases so the bot can evolve without "
                "manual SQL work every time a feature is added."
            ),
            p(
                "The database is not only storage; it is the bot's memory. It knows the active request wave, scheduled openings, "
                "submitted users and level IDs, request edit history, validation cache, weekly claims, weekly sessions, weekly "
                "reviews, activity counts, tickets, transcript pointers, help submissions, cooldowns, daily stats, and impact snapshots."
            ),
            diagram(
                "Persistence Safety Model",
                [
                    "Render Persistent Disk path",
                    "SQLite bot memory",
                    "Scheduled zipped backup",
                    "Discord backup channel",
                    "Impact and trend exports",
                ],
                "The primary durable copy is the mounted SQLite file; backup attachments and exports provide recovery evidence.",
            ),
            h2("Render Storage Rule"),
            p(
                "On Render, the project source and cache can be wiped by redeploys or cache clears. Avenue Guard therefore resolves "
                "its SQLite path from AVENUE_GUARD_DB_PATH first, then database.path in config.json, then an auto-detected "
                "Render Persistent Disk path at /var/data/avenue-guard/bot.db, and only then the local fallback. For production, "
                "mount a Render Persistent Disk at /var/data or point AVENUE_GUARD_DB_PATH at another durable path."
            ),
            p(
                "The /bot storage command checks the running path and the latest backup record. The /bot backup command creates a zipped "
                "copy immediately, and the background backup loop posts scheduled copies to the configured backup channel. If no persistent "
                "path is writable, the bot now starts with a local fallback and warns clearly, but that fallback should be treated as temporary."
            ),
            table(
                ["Table", "Purpose"],
                [
                    ["activity_counts", "Weekly message totals per user."],
                    ["weekly_claims / weekly_sessions / weekly_dm_log", "Weekly reward workflow state and audit history."],
                    ["weekly_request_reviews", "Weekly submitted request review messages and results."],
                    ["tickets / ticket_sequences", "Ticket channels, users, status, IDs, satisfaction, and closure state."],
                    ["ticket_transcripts / transcript_requests", "Saved transcript locations and member transcript request decisions."],
                    ["help_submissions / help_sessions / help_cooldowns", "DM help flows, appeal/report/bug submissions, and rate limits."],
                    ["level_request_state", "Current live request state, wave ID, limits, timers, and request button pointer."],
                    ["level_request_submissions", "Per-wave submitted levels, requester, result, review, and embed data."],
                    ["level_request_edit_audit", "Before/after snapshots of request edits."],
                    ["gd_level_validation_cache", "Cached validation results from external GD providers."],
                    ["daily_stats / weekly_recaps", "Operational telemetry snapshots and private recap history."],
                    ["impact_snapshots", "Persistent impact report payloads produced by /bot impact."],
                    ["database_backups", "Backup timestamps, channels, message IDs, sizes, reasons, and filenames."],
                ],
            ),
            callout(
                "Persistence rule",
                "If a workflow can span a restart or must be auditable later, it belongs in SQLite rather than only in memory.",
            ),
            p(
                "The migration approach is intentionally additive. New columns are added if missing, older tables are normalized "
                "when their shape changes, and indexes are created for common lookups. This lets the live bot keep its history "
                "while gaining new features such as ticket opening message IDs, request edit audit entries, scheduled opening "
                "messages, weekly review data, and impact snapshots."
            ),
        ],
    ),
    (
        "Cog Architecture",
        [
            p(
                "Each cog owns a feature family. This keeps the code understandable even though the bot is large. The cogs do "
                "interact with each other, but mostly through named methods and shared database/config utilities. CommandsCog "
                "is the command hub, while the workflow cogs handle the long-running state machines."
            ),
            table(
                ["Cog", "Primary Responsibility"],
                [
                    ["ModCog", "Proof-channel restrictions and role-triggered DMs."],
                    ["TrackingCog", "Weekly activity tracking, weekly request reward DMs, anti-farm checks, and weekly request recording."],
                    ["HelpCog", "DM help dashboard, FAQ, appeal/report/bug submissions, tickets, transcripts, and satisfaction."],
                    ["MessageResponsesCog", "Configurable message-triggered auto-responses."],
                    ["StickyCog", "Sticky messages, forum first-message reminders, and required-word thread deletion/logging."],
                    ["RequestLevelsCog", "Live request waves, scheduled openings, validation, modals, edit windows, reviews, results, summaries, and repairs."],
                    ["CommandsCog", "Admin, tracking, ticket, forum, request, server icon, fun, diagnostics, and impact commands."],
                    ["BackgroundCog", "Daily stats, summaries, rotating status, server icon rotation, and background persistence."],
                ],
            ),
            h2("Interaction Flow"),
            p(
                "Persistent views in utils.views act like switchboards. A button custom ID identifies the action, the view asks "
                "Discord for the relevant cog, and then the cog handles the real logic. This means the visible component can "
                "remain tiny while the business rules stay in the owning cog."
            ),
            p(
                "For example, the request button view only knows that a member clicked Request your level. RequestLevelsCog then "
                "checks whether requests are open, whether the member has the required role, whether they already submitted in "
                "the wave, whether the button should open an edit flow, and whether a modal should appear."
            ),
            p(
                "This separation is especially useful for persistent components. A button may be clicked long after the message "
                "was created, so the button itself should not carry fragile state. It carries a stable custom ID, and the owning "
                "cog retrieves current state from config and SQLite at click time."
            ),
        ],
    ),
    (
        "Permissions And Security Model",
        [
            p(
                "The bot combines Discord permissions, configured role IDs, and command-level checks. Public commands are kept "
                "narrow, mod commands require the configured mod role or configured permission policy, and admin commands require "
                "one of the configured admin/owner roles. Sensitive interactions also check the user before editing dashboards or panels."
            ),
            p(
                "Avenue Guard also avoids unsafe mention behavior in staff logs and bot-generated messages where possible. The "
                "no_mentions helper prevents accidental mass pings in logs and auto-responses. Where pings are intentional, such "
                "as the default request-open announcement, the behavior is explicit and configurable."
            ),
            bullets(
                [
                    "Guild restriction prevents the bot from operating outside the configured server.",
                    "Admin commands are role-gated even if their command descriptions do not visually say so.",
                    "Mod workflows check staff role or manage-guild policy before ticket/status operations.",
                    "Request reviewer controls are limited by access to the review channel and configured reviewer roles where staff filters apply.",
                    "Auto-response output is length-limited and mass mentions are blocked.",
                    "External validation has per-user rate limits and provider backoff to reduce abuse and failure cascades.",
                ]
            ),
            callout(
                "Security posture",
                "The bot is not a bank-grade security system, but it uses practical Discord safety controls: role gates, guild gates, safe mentions, cooldowns, audit logs, and recovery commands.",
            ),
            p(
                "Data safety is treated pragmatically. The bot stores IDs, message pointers, submitted text, review text, ticket "
                "metadata, and transcript pointers because those are necessary for accountability. It avoids storing secrets in "
                "the database and does not require Google credentials for impact reporting. Sensitive records should still be "
                "protected by keeping the database on trusted storage and limiting staff-log channel access."
            ),
        ],
    ),
    (
        "Moderation Guardrails",
        [
            p(
                "The guardrail layer handles repetitive moderation actions that should not depend on a staff member being online. "
                "The proof-channel restriction watches a configured channel. If a non-whitelisted member posts there, the bot "
                "deletes the message and applies the configured restriction role. If they add a reaction there, it removes the "
                "reaction and applies the same restriction role."
            ),
            p(
                "Role-triggered DMs are another guardrail. When a member gains a watched role, the bot sends a configured DM that "
                "explains what changed and how to appeal or contact staff. This turns silent role changes into explainable actions."
            ),
            h2("Forum And Sticky Reminders"),
            p(
                "Sticky messages keep important instructions visible at the bottom of busy text channels. The bot debounces sticky "
                "updates, deletes the old sticky message, and posts a fresh one after the configured delay. Forum first-message "
                "reminders post an embed in new forum threads, with tag-specific templates when configured."
            ),
            p(
                "Required-word enforcement is designed for forum formats that must include a specific word. The bot checks thread "
                "title/body text, supports contains, whole word, and regex modes, sends a configurable DM to the thread owner, "
                "deletes the thread after the configured delay, and logs the deletion with the author and thread context."
            ),
            callout(
                "Why this exists",
                "Forum reminders are gentle guidance; required-word enforcement is the hard stop for posts that ignore a required format.",
            ),
        ],
    ),
    (
        "Live Request System",
        [
            p(
                "The live request system is the bot's most involved workflow. It starts with a persistent request button embed in "
                "the configured request channel. Staff can refresh or recreate that embed with /refresh-request-button. Admins can "
                "open requests immediately, close them manually, or schedule openings for later."
            ),
            diagram(
                "Live Request Wave",
                [
                    "Open or scheduled opening",
                    "Member presses request button",
                    "Requirements and validation",
                    "Review queue embed",
                    "Send, Reject, or Other result",
                    "Wave summary update",
                ],
                "Submission count increases only after a valid modal succeeds, not when the button is pressed.",
            ),
            p(
                "A wave begins whenever requests open. A wave can be unlimited, limited by successful submission count, limited by "
                "time, or limited by both. If both count and time are defined, the count limit wins. A request only counts after a "
                "valid modal submission succeeds. Clicking the button or opening the form does not consume a slot."
            ),
            h2("Per-Wave Rules"),
            bullets(
                [
                    "One user can submit one live request per wave.",
                    "One level ID can be submitted once per wave.",
                    "Per-user and per-level duplicate tracking resets when a new wave starts.",
                    "Requests can be edited until the wave closes plus the configured grace period.",
                    "The wave summary is updated as reviews happen so staff can see remaining workload.",
                ]
            ),
            h2("Request Types"),
            p(
                "Request waves can optionally define a type, such as needs showcase, only demons, only platformers, only classic, "
                "classic non-demons, platformer non-demons, or long/XL levels. These types are enforced after validation when the "
                "bot has enough GD metadata to reason about difficulty, platformer status, and length."
            ),
            p(
                "Opening announcements are configurable. If no custom message is provided, the bot uses the default request role "
                "ping and inserts a human-readable condition summary. Scheduled openings can also store a custom opening message."
            ),
            p(
                "Scheduled openings are deliberately managed as records instead of timers only in memory. Admins can list, edit, "
                "delete, refresh, or open them immediately. If the bot restarts before the scheduled time, the pending opening "
                "still exists in SQLite and the scheduled-opening loop can act on it when the bot comes back."
            ),
        ],
    ),
    (
        "Request Validation And GD Metadata",
        [
            p(
                "Validation protects the request queue from bad level IDs and gives reviewers more context. The bot checks level "
                "IDs before accepting a modal: IDs must be 7 to 9 digits, showcase links must be URLs, and missing levels can be "
                "auto-rejected when enabled providers confidently agree that the ID does not exist."
            ),
            p(
                "The validation layer uses two providers: GDBrowser and the direct GD/Boomlings endpoint. Results are combined into "
                "one normalized payload that can include level name, creator, difficulty, length, stars, rated status, featured/epic "
                "flags, demon status, and platformer status. The result is cached in SQLite to keep repeated checks fast and to avoid "
                "hammering external services."
            ),
            table(
                ["Validation Output", "How It Is Used"],
                [
                    ["exists", "Blocks confidently missing IDs before they enter the review queue."],
                    ["rated", "Warns reviewers that a level may already be rated."],
                    ["demon/platformer", "Requires a showcase URL automatically."],
                    ["difficulty/length/stars", "Adds clean GD info to request embeds."],
                    ["provider disagreement", "Warns staff instead of hiding uncertainty."],
                    ["cache expiry", "Lets repair or new submissions refresh stale warnings later."],
                ],
            ),
            callout(
                "Validation principle",
                "The bot is strict only when the evidence is strong. When providers disagree or fail, the bot surfaces a warning instead of pretending certainty.",
            ),
            p(
                "Validation also feeds presentation. The request embed can show compact GD info without overcrowding the request: "
                "difficulty, length, stars/rated status, flags, creator, and provider warnings can be collapsed into clean fields. "
                "That means reviewers spend less time opening external pages just to understand what kind of level they are judging."
            ),
        ],
    ),
    (
        "Review Workflow",
        [
            p(
                "After a successful request submission, the bot sends a configurable embed to the level_requested channel. The "
                "embed includes requester, level ID, level name, creators, showcase, notes, GD info, validation warning, duplicate "
                "history warning, edit trail count, and wave information. The same view provides Send, Reject, and Other buttons."
            ),
            p(
                "Send and Reject open a review modal with an optional review field. Once submitted, the original request embed is "
                "edited into its final state, the result color changes, the reviewer is recorded, the result embed is posted to the "
                "sent or rejected channel, the requester is pinged there, and all buttons on the original request are disabled."
            ),
            p(
                "The Other button offers fixed reasons: level does not exist, stolen level, and already rated. These are treated like "
                "not-sent results and notify the requester through the rejected channel. This keeps special rejection reasons structured "
                "rather than buried in arbitrary review text."
            ),
            h2("Wave Summary"),
            p(
                "When a wave exists, the bot maintains a summary embed in level_requested. It shows total requested, reviewed count, "
                "sent count, not-sent count, percentages, remaining reviews, not-sent breakdown, and reviewer stats. This is the staff "
                "dashboard for the wave, and it updates each time a request is reviewed."
            ),
            h2("Repair"),
            p(
                "/requests repair exists because Discord messages can be deleted, embeds can go stale, validation warnings can expire, "
                "and reviewed messages should stay locked. The repair command refreshes the request button, rebuilds summaries, recreates "
                "missing pending request messages, refreshes validation warnings, and disables buttons on reviewed messages."
            ),
            p(
                "Review actions are designed to be idempotent from a staff perspective. The bot checks the original request row, verifies "
                "that it is still pending, confirms the result channel, edits the original embed, writes the review fields, sends the "
                "final notification, and disables buttons. This reduces the chance that two reviewers can accidentally process the same "
                "request twice."
            ),
        ],
    ),
    (
        "Weekly Activity And Rewards",
        [
            p(
                "TrackingCog counts eligible member messages by week. It skips excluded channels and roles, uses a cooldown to avoid "
                "overcounting rapid-fire messages, buffers writes to reduce SQLite load, and applies anti-farm checks before messages "
                "are added to the weekly leaderboard."
            ),
            p(
                "At reward time, the bot contacts configured winners through DM. A member can claim, decline, time out, or receive a "
                "reminder. The weekly claim tables and logs store who was contacted, what happened, and which user should be offered "
                "the next slot if someone declines or times out. Admins can disable and re-enable the automatic reward for the current week."
            ),
            p(
                "Weekly request submissions use the same Send, Reject, and Other review workflow as live requests, but they are not part "
                "of a live request wave. This means staff review behavior stays consistent while wave-specific limits and summaries remain "
                "clean."
            ),
            h2("Streaks And Anti-Farm"),
            p(
                "Weekly streaks reward members who repeatedly place in the configured top rank band. Anti-farm detection watches for "
                "repeated low-effort messages and logs suspicious patterns instead of letting them inflate weekly counts. The result is a "
                "leaderboard that is harder to game and more useful for community reward decisions."
            ),
            p(
                "Manual force-DM exists for operational exceptions. Admins can send the weekly request DM to a member even if normal "
                "tracking would exclude them or the automatic reward is disabled for the week. The result is logged so manual overrides "
                "remain visible to future staff."
            ),
        ],
    ),
    (
        "Help, Tickets, And Staff Support",
        [
            p(
                "The DM help system starts from a dashboard. Members can see active ticket status, weekly activity status, current request "
                "state, recent help submissions, and cooldowns. The menu hides the option the user is already viewing, cleans up previous "
                "screens when possible, and supports Back, Cancel, and Start Over controls."
            ),
            diagram(
                "Help And Ticket Flow",
                [
                    "DM dashboard",
                    "FAQ suggestions",
                    "Submission or ticket",
                    "Staff log or private channel",
                    "Reply, transcript, or satisfaction",
                ],
                "The user experience stays private while staff still get auditable records.",
            ),
            p(
                "FAQ search and auto-suggestions are meant to reduce unnecessary tickets. Before opening a staff ticket, the bot can show "
                "relevant FAQ entries so common questions are solved privately. If the user still needs help, they can open a routed private "
                "ticket channel by topic."
            ),
            h2("Submission Workflows"),
            p(
                "Appeals, user reports, bot issue reports, and transcript requests use tracked submissions. The bot stores a code, keeps "
                "attachment links, shows a preview before submission, posts a structured staff log embed, and lets staff reply to a log "
                "message to relay a response back to the submitter by DM."
            ),
            h2("Tickets"),
            p(
                "Tickets use atomic ticket IDs, private channels, status tags, inactivity scans, close prompts, transcripts, transcript "
                "search, and satisfaction prompts. The opening message is kept in sync when staff or users reply, when a staff member changes "
                "status, and when the ticket closes. Before deletion, the bot saves a transcript and records where the transcript was posted."
            ),
            p(
                "The help system is intentionally private-first. It gives members a place to ask for help without escalating every issue "
                "into a public channel, but it still creates staff-visible logs when something becomes an official submission. This balances "
                "member comfort with staff accountability."
            ),
        ],
    ),
    (
        "Background Telemetry",
        [
            p(
                "BackgroundCog is the bot's measurement layer. It listens for messages, edits, deletes, reactions, joins, leaves, bans, "
                "unbans, boosts, voice state changes, command completions, and command errors. These events are accumulated into daily "
                "snapshots and persisted in daily_stats."
            ),
            diagram(
                "Telemetry Flow",
                [
                    "Discord activity",
                    "Daily counters",
                    "daily_stats rows",
                    "Daily or weekly embeds",
                    "Impact forecast exports",
                ],
                "Operational telemetry is useful for trends, but it should be described as tracked data rather than absolute community reality.",
            ),
            p(
                "The daily summary embed turns raw counters into something readable: message totals, day-over-day movement, active members, "
                "active channels, joins/leaves, moderation signals, voice time, command success rate, top channels, top members, and top commands. "
                "Weekly recaps summarize longer-term activity, request, review, streak, and anti-farm patterns."
            ),
            h2("Presence And Icon Rotation"),
            p(
                "The bot can rotate its Discord status using placeholders such as members, online count, weekly messages, current top member, "
                "open tickets, and today's messages. It can also rotate the server icon through configured image URLs in disabled, linear, or "
                "random modes. The icon rotation code downloads images, checks that they look like supported image bytes, remembers failures, "
                "and stores current index/state back into config.json."
            ),
            p(
                "Daily stats are useful but should be read as operational telemetry, not perfect analytics. They depend on bot uptime, enabled "
                "intents, cache visibility, and events the bot can observe. That is why impact reports label large totals as tracked events "
                "rather than claiming to represent every possible interaction in the community."
            ),
        ],
    ),
    (
        "Admin Tools And Diagnostics",
        [
            p(
                "CommandsCog exposes most operator-facing slash commands. It includes tracking commands, ticket commands, forum required-word "
                "management, request review filters, request history, request repair, server icon controls, fun commands, and bot diagnostics."
            ),
            p(
                "The admin dashboard is a button-driven status view. It gathers system health, request state, tracking state, icon rotation, "
                "config issues, and repair suggestions into one embed. This reduces the need for scattered health commands while still keeping "
                "older commands available for direct checks."
            ),
            table(
                ["Diagnostic", "Purpose"],
                [
                    ["/bot dashboard", "Interactive overview of system health, config, and repair tips."],
                    ["/bot health", "Compact live health report."],
                    ["/bot config_check", "Checks configured channels, roles, templates, and response rules."],
                    ["/bot doctor", "Deeper permission and system diagnostics."],
                    ["/requests repair", "Repairs request system messages, validation warnings, summaries, and locks."],
                    ["/bot impact", "Owner-only impact and forecast exports with Markdown, CSV, trend CSV, breakdown CSV, and JSON."],
                    ["/bot backup", "Creates a zipped SQLite backup and posts it to the configured backup channel."],
                    ["/bot storage", "Shows active database path, persistence warning, backup channel, interval, and latest backup."],
                ],
            ),
            callout(
                "Operator principle",
                "When a feature can fail because a Discord message, channel, permission, or config value changed, the bot should expose a command that explains or repairs it.",
            ),
        ],
    ),
    (
        "Impact Reporting",
        [
            p(
                "The owner-only /bot impact command turns the bot's persistent state into a quantifiable impact report. It collects current "
                "server size, unique members touched by tracked workflows, tracked interaction events, support/help volume, ticket volume, "
                "transcripts, request totals, review rates, weekly reward activity, command usage, voice minutes, anti-farm events, and "
                "summary history."
            ),
            diagram(
                "Impact Reporting Pipeline",
                [
                    "SQLite workflow tables",
                    "Daily trend rows",
                    "Forecast and backlog metrics",
                    "Markdown, CSV, JSON files",
                    "Discord evidence channel",
                ],
                "The report is both human-readable and spreadsheet-ready so it can support CV evidence and operational planning.",
            ),
            p(
                "The command posts a Markdown report, summary CSV, daily trend CSV, breakdown CSV, and raw JSON file to the configured "
                "impact report channel, then stores the same report payload in the impact_snapshots database table. The Markdown file is "
                "human-readable and CV-friendly. The CSV files can be imported into Google Sheets for charts, portfolio evidence, forecasting, "
                "or regular impact tracking."
            ),
            p(
                "The report now includes a simple forecast model. It compares the last seven days with the previous seven days, projects "
                "the next seven days from that movement, labels the engagement signal, and highlights review backlog or command error risk. "
                "This should be read as an operations forecast, not a perfect prediction."
            ),
            h2("Why This Is Defensible For A CV"),
            p(
                "The report uses numbers the bot actually tracks. Instead of claiming vague community influence, it produces concrete "
                "figures such as members reached, support items handled, level requests coordinated, tracked events, tickets resolved, and "
                "review throughput. This makes the result useful for a CV because it describes operational impact in measurable terms."
            ),
            p(
                "For best evidence, run /bot impact on a recurring cadence such as monthly or before major application updates. Keep the "
                "posted Discord attachments, and import the CSV files into a spreadsheet when you want trend charts. The database snapshot is useful "
                "for bot-side history, while the Discord attachment gives you a durable, shareable artifact."
            ),
            callout(
                "Example CV wording",
                "Built and maintained Avenue Guard, a Discord operations bot supporting a multi-thousand-member community, coordinating "
                "level request workflows, staff tickets, weekly rewards, help flows, moderation guardrails, and persistent impact reporting.",
            ),
        ],
    ),
    (
        "External Services And Dependencies",
        [
            p(
                "Avenue Guard relies on Discord as its primary platform, py-cord as its Discord framework, aiohttp for asynchronous HTTP, "
                "SQLite for persistence, and optional hosted infrastructure for runtime availability. Most data remains local to the bot's "
                "database and Discord channels."
            ),
            p(
                "The Geometry Dash validation feature uses GDBrowser and the GD/Boomlings endpoint. These services can fail, disagree, rate "
                "limit, or return unexpected payloads. The bot handles that by normalizing provider responses, caching results, surfacing "
                "warnings, and backing off providers that fail repeatedly."
            ),
            h2("Google Sheets Consideration"),
            p(
                "The bot now exports multiple CSV impact files. That is the safest immediate bridge to Google Sheets because it does not require "
                "storing Google credentials in the bot. If a future service account or Google Drive integration is added, the same metrics "
                "payload can be uploaded automatically. Until then, the summary, trend, and breakdown CSV files are designed to import cleanly "
                "into a spreadsheet."
            ),
        ],
    ),
    (
        "Failure Recovery",
        [
            p(
                "Avenue Guard assumes that Discord state can drift. A message can be deleted, a channel can be moved, a role can be missing, "
                "a permission can change, a provider can fail, or a database can be older than the current code. Recovery is therefore built "
                "into migrations, diagnostics, admin logs, repair commands, and cautious external validation."
            ),
            bullets(
                [
                    "Database migration creates missing tables and columns on startup.",
                    "Global command and event error handlers log failures instead of silently swallowing them.",
                    "Request repair can rebuild missing request messages and relock reviewed embeds.",
                    "Ticket close restores status if transcript/close fails partway through.",
                    "Weekly request recording failures are logged and do not silently mark claims as successful.",
                    "Icon rotation remembers last errors and avoids changing too frequently.",
                    "Impact reports persist both a DB payload and Discord attachments when the report channel is configured.",
                    "Scheduled database backups post zipped SQLite copies to Discord when the backup channel is configured.",
                ]
            ),
            callout(
                "Recovery philosophy",
                "The bot does not need to be impossible to break. It needs to fail visibly, preserve state, and provide a clear path back to a working condition.",
            ),
            p(
                "In practice, most failures fall into a few categories: config points at a missing channel, the bot lacks a permission, a "
                "message was deleted, an external provider failed, a user disabled DMs, or a deploy restarted the process mid-workflow. The "
                "bot's current recovery tools are aimed at exactly those categories."
            ),
        ],
    ),
    (
        "Maintenance And Testing",
        [
            p(
                "The main test guide is TEST_CHECKLIST.md. It is intentionally server-side because many behaviors require Discord state: "
                "roles, channels, messages, DMs, buttons, modals, slash command permissions, forum threads, scheduled tasks, and external "
                "request validation."
            ),
            p(
                "Code-level checks still matter. The project should compile cleanly, config.json should parse, and database migrations should "
                "run against a temporary database. For risky changes, test the real Discord workflow with a staff account and a non-staff account."
            ),
            h2("Recommended Maintenance Routine"),
            numbered(
                [
                    "Run a syntax and config check before deploying.",
                    "Run /bot dashboard after deploying to catch missing roles, channels, or permissions.",
                    "Use /requests repair after request-template, validation, or message-state changes.",
                    "Run /bot storage after deploying to confirm the database path is persistent.",
                    "Run /bot backup after first deploy and before major migrations.",
                    "Run /bot impact periodically and keep the CSV files for trend tracking.",
                    "Update this manual when new feature families are added.",
                ]
            ),
        ],
    ),
    (
        "Appendix: Command Families",
        [
            p(
                "The bot exposes commands by family so staff can discover tools without memorizing every implementation detail. Command "
                "descriptions are kept clean; role restrictions are enforced by code instead of being advertised awkwardly in every description."
            ),
            table(
                ["Family", "Commands"],
                [
                    ["Tracking", "/tracking top, /tracking me, /tracking reset, /tracking force_dm, /tracking disable_reward, /tracking enable_reward"],
                    ["Requests", "/refresh-request-button, /open-requests, /close-requests, /requests-are, /edit-request, /pending-openings, /requests pending, /requests history, /requests repair"],
                    ["Tickets", "/ticket close, /ticket status, /ticket transcripts"],
                    ["Forum", "/forum required_word"],
                    ["Bot/admin", "/bot dashboard, /bot health, /bot config_check, /bot doctor, /bot impact, /bot backup, /bot storage, /resync, /restart"],
                    ["Server icon", "/server_icon status, /server_icon mode, /server_icon add, /server_icon replace, /server_icon remove, /server_icon set, /server_icon next"],
                    ["Fun", "/dance, /rock-paper-scissors, /gambling"],
                ],
            ),
            h2("Operating Rule Of Thumb"),
            p(
                "Use public commands for member self-service, mod commands for ticket and forum operations, admin commands for stateful or "
                "config-affecting actions, and repair/doctor commands whenever Discord state no longer matches the database."
            ),
        ],
    ),
    (
        "Appendix: Data And Evidence",
        [
            p(
                "Avenue Guard's strongest evidence is the data it already generates. Tickets, transcripts, help submissions, request waves, "
                "weekly claims, daily stats, and impact reports can show how much community work the bot has handled. The important thing is "
                "to use labels that match what is measured."
            ),
            table(
                ["Metric Label", "Source", "Good Use"],
                [
                    ["Current server members", "Discord guild member count", "Shows the size of the community the bot supports."],
                    ["Unique members touched", "Union of tracked workflow user IDs", "Shows historical reach across bot workflows."],
                    ["Tracked interaction events", "Messages, commands, requests, tickets, help, DMs, reviews, transcripts, and safety logs", "Shows operational throughput, not every human action in the server."],
                    ["Support/help items", "Tickets, help submissions, transcript requests", "Shows staff-support workload handled by the bot."],
                    ["Level requests coordinated", "Live requests plus weekly request reviews", "Shows request-program volume."],
                    ["Review rate", "Reviewed requests divided by total requests", "Shows staff queue completion."],
                ],
            ),
            p(
                "For a CV, the safest wording combines a clear build claim with measured impact. For example: Built Avenue Guard, a Discord "
                "operations bot for GD Avenue that automates request waves, weekly rewards, tickets, help workflows, moderation guardrails, "
                "and analytics, with persistent reports quantifying member reach, support volume, request throughput, and staff review outcomes."
            ),
        ],
    ),
]


def _source_file_inventory_rows() -> list[list[str]]:
    files = [
        "main.py",
        "cogs/Background.py",
        "cogs/Commands.py",
        "cogs/Help.py",
        "cogs/MessageResponses.py",
        "cogs/Mod.py",
        "cogs/RequestLevels.py",
        "cogs/Sticky.py",
        "cogs/Tracking.py",
        "utils/config.py",
        "utils/db.py",
        "utils/errors.py",
        "utils/gd_validation.py",
        "utils/server_icons.py",
        "utils/views.py",
    ]
    rows: list[list[str]] = []
    for rel in files:
        path = ROOT / rel
        try:
            text = path.read_text(encoding="utf-8")
            parsed = ast.parse(text)
            classes = sum(isinstance(node, ast.ClassDef) for node in ast.walk(parsed))
            functions = sum(isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)) for node in ast.walk(parsed))
            listeners = len(re.findall(r"@commands\.Cog\.listener", text))
            tasks = len(re.findall(r"@tasks\.loop", text))
            rows.append([rel, str(len(text.splitlines())), f"{classes} classes / {functions} functions", f"{listeners} listeners / {tasks} loops"])
        except Exception as exc:
            rows.append([rel, "?", f"Could not parse: {type(exc).__name__}", ""])
    return rows


def _command_inventory_rows() -> list[list[str]]:
    rows: list[list[str]] = []
    files = ["cogs/Commands.py", "cogs/RequestLevels.py"]
    command_re = re.compile(r'(?:slash_command|command)\(name="([^"]+)",\s*description="([^"]*)"', re.M)
    direct_re = re.compile(r'@bot\.slash_command\(name="([^"]+)",\s*description="([^"]*)"', re.M)
    for rel in files:
        text = (ROOT / rel).read_text(encoding="utf-8")
        seen: set[str] = set()
        for name, desc in command_re.findall(text) + direct_re.findall(text):
            key = f"{rel}:{name}:{desc}"
            if key in seen:
                continue
            seen.add(key)
            rows.append([f"/{name}", rel, desc])
    rows.sort(key=lambda row: (row[1], row[0]))
    return rows


def _database_table_rows() -> list[list[str]]:
    text = (ROOT / "utils/db.py").read_text(encoding="utf-8")
    names = sorted(set(re.findall(r"CREATE TABLE IF NOT EXISTS\s+([A-Za-z0-9_]+)", text)))
    purpose = {
        "activity_counts": "Weekly activity totals.",
        "activity_last_counted": "Per-user cooldown memory for activity counting.",
        "weekly_claims": "Weekly reward contact and claim status.",
        "weekly_sessions": "Active weekly DM claim sessions.",
        "weekly_dm_log": "Weekly workflow audit events.",
        "weekly_reminders": "Reminder delivery memory.",
        "weekly_runs": "Scheduler idempotency for weekly jobs.",
        "weekly_reward_disabled": "Per-week switch for automatic reward delivery.",
        "weekly_streaks": "Top-member streak tracking.",
        "weekly_recaps": "Private weekly recap message history.",
        "anti_farm_events": "Skipped low-effort activity events.",
        "weekly_request_reviews": "Weekly request review queue rows.",
        "tickets": "Ticket channel, owner, status, ID, and satisfaction state.",
        "ticket_sequences": "Atomic ticket ID counters by guild.",
        "ticket_transcripts": "Transcript log message pointers.",
        "ticket_cooldowns": "Ticket creation cooldowns.",
        "sticky_state": "Last sticky message per channel.",
        "help_sessions": "Current DM help flow stage.",
        "help_cooldowns": "Help action rate limits.",
        "help_submissions": "Appeals, reports, and bot issue records.",
        "transcript_requests": "Member transcript approval requests.",
        "rps_streaks": "Rock-paper-scissors win streaks.",
        "level_request_state": "Current live request wave state.",
        "level_request_submissions": "Live request submissions and reviews.",
        "level_request_wave_summaries": "Live wave summary message pointers.",
        "level_request_scheduled_openings": "Future request openings.",
        "level_request_edit_audit": "Before/after request edit snapshots.",
        "gd_level_validation_cache": "Cached GD provider validation payloads.",
        "daily_stats": "Daily telemetry payloads.",
        "impact_snapshots": "Owner impact report payload history.",
        "database_backups": "Posted backup metadata.",
    }
    return [[name, purpose.get(name, "Persistent workflow state.")] for name in names]


def _private_deep_dive_chapters() -> list[tuple[str, list[tuple]]]:
    return [
        (
            "Private Orientation: How To Study The Bot",
            [
                p(
                    "This private section is written for you as the builder-owner of Avenue Guard. The goal is not only to know "
                    "which commands exist. The goal is to understand how a Discord event becomes code, how code turns into stored "
                    "state, how that state survives a restart, and how the bot repairs visible Discord messages when they drift."
                ),
                p(
                    "A useful way to study the project is to read it in layers: main.py starts the system, utils provide shared "
                    "rules, cogs own workflows, config.json controls server-specific behavior, and SQLite remembers anything that "
                    "matters after a restart. When you get lost, ask: who owns this event, where is the state stored, and what "
                    "Discord object does the user see?"
                ),
                diagram(
                    "Code Reading Map",
                    [
                        "main.py boots",
                        "utils define shared rules",
                        "cogs own workflows",
                        "config shapes behavior",
                        "SQLite preserves state",
                        "Discord shows results",
                    ],
                    "This is the mental route for almost every feature in the bot.",
                ),
                table(
                    ["Question", "Where To Look First"],
                    [
                        ["Why did the bot start or fail?", "main.py and utils/db.py"],
                        ["Why did a command answer this way?", "cogs/Commands.py or cogs/RequestLevels.py command method"],
                        ["Why did a button do something after restart?", "utils/views.py custom ID and the owning cog handler"],
                        ["Why did a request enter or skip the queue?", "RequestLevelsCog validation, requirements, and submission lock"],
                        ["Why did a weekly winner get contacted?", "TrackingCog weekly loop, weekly_claims, weekly_sessions"],
                        ["Why did a ticket status change?", "HelpCog on_message and ticket status helpers"],
                        ["Why did a report number appear in impact data?", "CommandsCog _collect_impact_metrics"],
                    ],
                ),
                callout(
                    "The important pattern",
                    "Avenue Guard is event-driven, but its serious workflows are state-driven. The event starts the logic; the database decides what is true.",
                ),
            ],
        ),
        (
            "Source Inventory",
            [
                p(
                    "This chapter is generated from the current source files. It gives you a quick structural map before the deeper "
                    "walkthroughs. Line counts are not a quality metric by themselves, but they reveal where most of the bot's complexity lives."
                ),
                table(["File", "Lines", "Shape", "Discord Hooks"], _source_file_inventory_rows()),
                p(
                    "The largest files are large because they own full workflows, not because they only hold utility helpers. "
                    "RequestLevels.py owns a state machine with modals, validation, buttons, scheduled openings, review actions, and repairs. "
                    "Commands.py owns the command surface and cross-system diagnostics. Help.py owns the DM and ticket state machines."
                ),
                h2("How To Use This Inventory"),
                p(
                    "When debugging, avoid starting from the biggest file and scrolling randomly. Start from the user action. If it is a "
                    "slash command, search the command name. If it is a button, search the custom ID in utils/views.py and follow the handler. "
                    "If it is a background action, search the task loop name or the database table it changes."
                ),
            ],
        ),
        (
            "Command Surface Inventory",
            [
                p(
                    "This table is extracted from command registrations. The bot has direct commands and grouped commands. Some legacy or "
                    "programmatic commands appear by short name here, while their actual Discord path may include a group such as /bot, "
                    "/tracking, /ticket, /requests, /forum, or /server_icon."
                ),
                table(["Command", "Registered In", "Description"], _command_inventory_rows()),
                callout(
                    "Descriptions are intentionally clean",
                    "The code enforces permissions. The command descriptions do not need to carry visual labels like admin-only or owner-only.",
                ),
            ],
        ),
        (
            "Database Table Inventory",
            [
                p(
                    "This table is extracted from utils/db.py. It is one of the most useful quick-reference sections because almost every "
                    "serious Avenue Guard behavior has a table behind it."
                ),
                table(["Table", "Purpose"], _database_table_rows()),
                p(
                    "The table names are grouped by feature family. activity_* and weekly_* belong to tracking. ticket_* and help_* belong "
                    "to support. level_request_* and gd_level_validation_cache belong to live requests. daily_stats, impact_snapshots, and "
                    "database_backups belong to measurement and durability."
                ),
            ],
        ),
        (
            "Startup Code Walkthrough",
            [
                p(
                    "main.py is the bot's entry point. The most important design choice here is that startup resolves a usable database path "
                    "before creating the Database wrapper. That protects production from accidentally using ephemeral source storage when a "
                    "persistent disk exists."
                ),
                source_code("Database path resolution", "main.py", 46, 74),
                p(
                    "The path resolver checks sources in priority order: environment variable, config.json, Render persistent disk candidate, "
                    "then local fallback. The key technical trick is the write probe: the bot does not assume a path works just because it looks "
                    "right. It tries to create the parent directory, writes a tiny test file, deletes it, and only then accepts the path."
                ),
                source_code("Bot creation and startup hooks", "main.py", 76, 138),
                p(
                    "create_bot wires the bot object, configuration, database, cogs, and on_ready behavior together. on_ready is where the "
                    "runtime becomes alive: database connection, guild validation, background loops, and persistent views are all started from there."
                ),
                source_code("Persistent view registration", "main.py", 167, 177),
                p(
                    "Persistent view registration is easy to underestimate. Discord button messages can outlive the Python process. Without "
                    "registering the views again after restart, users could click old buttons and Discord would not know which callback should run."
                ),
            ],
        ),
        (
            "Database Code Walkthrough",
            [
                p(
                    "The Database wrapper is small because it has one job: make SQLite safe enough for an async Discord bot. SQLite calls are "
                    "blocking, so the wrapper serializes access with an asyncio.Lock and runs the actual SQLite work inside asyncio.to_thread."
                ),
                source_code("Connection, WAL, migration, and backup", "utils/db.py", 14, 66),
                p(
                    "WAL mode helps SQLite handle concurrent readers while writes are happening. The lock still serializes bot-side operations, "
                    "which prevents two coroutine paths from sharing one cursor incorrectly. This is less glamorous than a bigger database, but it "
                    "fits a single-server bot well and keeps deployment simple."
                ),
                source_code("Atomic ticket sequence and query helpers", "utils/db.py", 935, 1003),
                p(
                    "The ticket ID function is the cleanest example of an atomic counter in this codebase. It reads and increments under the "
                    "same database lock, commits before returning, and stores the next value by guild. This prevents two tickets opened at nearly "
                    "the same time from receiving the same visible ticket number."
                ),
                callout(
                    "Atomic means indivisible",
                    "In this bot, atomic usually means one protected operation that cannot be interleaved with another coroutine halfway through. "
                    "The lock plus one database transaction gives that guarantee for counters and state transitions.",
                ),
            ],
        ),
        (
            "Config And Template Engine Walkthrough",
            [
                p(
                    "config.json is treated as a practical control plane. The Config class is intentionally simple: load JSON, expose typed "
                    "getters, and save atomically through a temporary file plus os.replace. This matters because a partially-written config file "
                    "could break startup."
                ),
                source_code("Typed config getters and atomic save", "utils/config.py", 13, 65),
                p(
                    "Many embeds use Python format_map with a SafeDict. If a template references a missing variable, the bot inserts an empty "
                    "string instead of crashing the workflow. That is why template validation is useful: SafeDict keeps the bot alive, while "
                    "config checks help you notice mistakes before users see blank fields."
                ),
                source_code("Request embed template renderer", "cogs/RequestLevels.py", 1024, 1064),
                p(
                    "The embed renderer is shared by live request submissions, reviewed request embeds, result notifications, and wave summaries. "
                    "It reads fields, footer, images, thumbnails, author info, color, title, and description from config. Workflow logic stays in "
                    "Python; presentation stays in config."
                ),
            ],
        ),
        (
            "Persistent Views Walkthrough",
            [
                p(
                    "utils/views.py is the component router. It stores stable custom IDs and very small button/select classes. The view should "
                    "not implement the business rules. It should only receive the click and call the owning cog."
                ),
                source_code("Request button and review button router", "utils/views.py", 149, 205),
                p(
                    "This is why a request review button can still work after a restart. The custom ID is stable, the view is registered on "
                    "startup, and the callback asks the live bot instance for RequestLevelsCog. The cog then reloads the real request state from SQLite."
                ),
                diagram(
                    "Persistent Button Dispatch",
                    [
                        "Discord button click",
                        "Stable custom_id",
                        "View callback",
                        "get_cog lookup",
                        "Database-backed handler",
                    ],
                    "The button identifies the action; the database identifies the current truth.",
                ),
            ],
        ),
        (
            "Live Request Code Walkthrough",
            [
                p(
                    "The request button is deceptively complex. On click, the bot checks whether the user already has a current-wave request. "
                    "If they do and the edit window is still open, the same button becomes an edit entry point. If not, it checks open/closed "
                    "state, required roles, banned role, first-time request role logic, and then opens the modal."
                ),
                source_code("Request button state gate", "cogs/RequestLevels.py", 2078, 2164),
                p(
                    "The submission handler uses a lock because request limits and duplicate checks must be consistent. Imagine a wave with one "
                    "slot left and two users submit at the same moment. Without the lock, both could pass the count check. With the lock, one "
                    "complete submission finishes before the next one evaluates the current state."
                ),
                source_code("Request form core transaction", "cogs/RequestLevels.py", 2166, 2296),
                p(
                    "Notice the order: validate local fields, defer the interaction, validate externally, enter the submit lock, reload current "
                    "state, check duplicate user and duplicate level ID, send the review embed, store the Discord message ID, then increment the "
                    "wave count. The request only counts after the staff queue message exists."
                ),
                source_code("Request edit audit trail", "cogs/RequestLevels.py", 2368, 2447),
                p(
                    "The edit path writes both the new data and an audit record. That lets reviewers know the request changed and lets you inspect "
                    "what changed later. The audit table stores old and new JSON snapshots because request form data is template-driven and may "
                    "gain fields over time."
                ),
            ],
        ),
        (
            "Validation Code Walkthrough",
            [
                p(
                    "Validation is split into two files. utils/gd_validation.py knows how to parse provider responses and combine them. "
                    "RequestLevelsCog decides when to call validation, cache it, rate-limit it, and turn the result into user-facing errors "
                    "or reviewer warnings."
                ),
                source_code("Provider result combiner", "utils/gd_validation.py", 197, 264),
                p(
                    "The combiner does not pretend providers are always perfect. It tracks existing results, missing results, failed providers, "
                    "disagreement, rating status, and whether a showcase appears required. This is why the bot can block confidently missing IDs "
                    "but only warn when a provider failed or disagreed."
                ),
                source_code("Cached provider lookup and circuit breaker use", "cogs/RequestLevels.py", 833, 912),
                source_code("External validation policy", "cogs/RequestLevels.py", 976, 1012),
                p(
                    "The cache is important for both speed and kindness to external services. The circuit breaker is a practical resilience feature: "
                    "if one provider fails repeatedly, the bot temporarily stops using it instead of letting every submission wait on a broken service."
                ),
            ],
        ),
        (
            "Request Review Code Walkthrough",
            [
                p(
                    "Review actions are shared between live wave requests and weekly request submissions. The handler first figures out whether "
                    "the clicked message belongs to level_request_submissions or weekly_request_reviews, then applies the same result logic."
                ),
                source_code("Review target lookup and button gate", "cogs/RequestLevels.py", 2482, 2521),
                source_code("Final review transaction", "cogs/RequestLevels.py", 2523, 2613),
                p(
                    "The review lock has the same purpose as the submit lock: one reviewer should win the state transition. The database status "
                    "must still be pending when the review is saved. After that, the original buttons are disabled so Discord's visible UI matches "
                    "the stored result."
                ),
                source_code("Wave summary variables and reviewer stats", "cogs/RequestLevels.py", 1138, 1226),
                p(
                    "The summary is generated from the database, not from memory. That means it can be rebuilt later and it stays correct even if "
                    "the bot restarts between reviews."
                ),
            ],
        ),
        (
            "Scheduled Openings Walkthrough",
            [
                p(
                    "Scheduled request openings are stored as rows, not just sleeping tasks. This is deliberate. If the bot restarts, a sleeping "
                    "task disappears, but the row in level_request_scheduled_openings remains. The scheduler loop can pick it up again when the bot is ready."
                ),
                source_code("Scheduling command branch", "cogs/RequestLevels.py", 1833, 1879),
                source_code("Pending opening edit/list branch", "cogs/RequestLevels.py", 1881, 1968),
                p(
                    "The command accepts immediate openings and scheduled openings through the same entry point. The when/day options create a "
                    "future row; leaving when empty opens immediately. This keeps the admin interface compact while the stored state remains explicit."
                ),
                callout(
                    "Why Discord timestamps are used",
                    "The bot stores Unix timestamps and displays Discord timestamp markup. Discord then renders the time in each viewer's client, "
                    "which avoids putting a timezone label in the frontend while still keeping scheduling internally consistent.",
                ),
            ],
        ),
        (
            "Tracking Code Walkthrough",
            [
                p(
                    "TrackingCog counts activity without writing to SQLite on every message. That would be slow and noisy. Instead, it keeps "
                    "small in-memory buffers and periodically flushes them with UPSERT statements. If the flush fails, it puts the counts back "
                    "into the buffer so they can be retried."
                ),
                source_code("Message counting gate", "cogs/Tracking.py", 608, 668),
                source_code("Buffered activity flush", "cogs/Tracking.py", 675, 710),
                p(
                    "The ON CONFLICT SQL is doing the increment atomically at database level: if a row already exists for that user and week, "
                    "count becomes count + excluded.count. That keeps weekly totals correct even though the bot flushes multiple messages together."
                ),
                source_code("Weekly job runner", "cogs/Tracking.py", 1176, 1233),
                p(
                    "The weekly job flushes pending activity first, ranks users, updates streaks, respects the weekly reward disabled switch, "
                    "contacts winners, writes the weekly_runs idempotency row, and creates a private recap. The weekly_runs table prevents the "
                    "same week from being processed repeatedly by the scheduler."
                ),
            ],
        ),
        (
            "Weekly Request Workflow Walkthrough",
            [
                p(
                    "Weekly request rewards happen in DMs. A winner receives a formatted request prompt, replies with the required fields, and "
                    "TrackingCog parses the message into the same shape that RequestLevelsCog understands for review embeds."
                ),
                source_code("Weekly DM request parser", "cogs/Tracking.py", 98, 168),
                source_code("Weekly request recording", "cogs/Tracking.py", 772, 868),
                p(
                    "The important bridge is LevelRequestReviewView. Weekly submissions are not part of a live wave, but they still use the same "
                    "Send, Reject, and Other buttons. The review row lives in weekly_request_reviews, and RequestLevelsCog's review finalizer handles it."
                ),
                source_code("Weekly contact state", "cogs/Tracking.py", 1235, 1297),
                p(
                    "force_dm is an intentional override path. Normal weekly rewards respect exclusions and the disabled switch; manual force-DM "
                    "can be used for exceptions and is logged so the override is visible later."
                ),
            ],
        ),
        (
            "Help And Ticket Code Walkthrough",
            [
                p(
                    "HelpCog is a state machine for DMs and ticket channels. The user's current help stage is stored in help_sessions. A typed "
                    "message or button action reads the stage, updates the session, and sends the next prompt."
                ),
                source_code("Help session message router", "cogs/Help.py", 1042, 1117),
                p(
                    "The preview step exists to prevent accidental submissions. For appeals, reports, and bot issues, the user can review the "
                    "embed, edit the last answer, cancel, or submit. The staff log only receives the item after the preview is confirmed."
                ),
                source_code("Help submission insert and staff log", "cogs/Help.py", 1001, 1040),
                source_code("Ticket creation", "cogs/Help.py", 1689, 1765),
                p(
                    "Ticket IDs come from the database counter, not from Discord channel IDs. That creates short human labels like T123 while "
                    "still preserving the real channel ID for lookups, transcript indexing, and closure."
                ),
                source_code("Ticket closure safety", "cogs/Help.py", 1850, 1949),
                p(
                    "Ticket close is cautious. It marks the ticket resolved, builds a transcript, posts the transcript to the log channel, indexes "
                    "the transcript, deletes the channel, and prompts satisfaction. If a dangerous middle step fails, it restores the previous status "
                    "instead of deleting the channel blindly."
                ),
            ],
        ),
        (
            "Background Telemetry Code Walkthrough",
            [
                p(
                    "BackgroundCog turns many Discord events into a daily payload. The dataclass keeps today's counters in memory, while "
                    "daily_stats stores snapshots so restart and impact reporting do not wipe the day."
                ),
                source_code("DailyStats shape", "cogs/Background.py", 61, 86),
                source_code("Daily stat persistence", "cogs/Background.py", 473, 515),
                p(
                    "The rollover logic is subtle because voice time can span midnight. The bot calculates minutes up to the boundary, persists "
                    "the old day, then starts a fresh day with current voice sessions carried forward from the guild state."
                ),
                source_code("Daily message listener", "cogs/Background.py", 588, 602),
                source_code("Daily report embed", "cogs/Background.py", 933, 1082),
                p(
                    "The summary embed is built from the stored counters plus derived values: net member movement, command success rate, average "
                    "messages per active member, top channels, top users, and top commands. This is the raw material for later impact reports."
                ),
            ],
        ),
        (
            "Forum And Sticky Code Walkthrough",
            [
                p(
                    "StickyCog has two different jobs that both involve keeping instructions visible: channel sticky messages and forum first "
                    "messages. It also enforces the required-word rule for forum threads."
                ),
                source_code("Sticky debounced repost", "cogs/Sticky.py", 120, 181),
                source_code("Required word detection", "cogs/Sticky.py", 276, 324),
                source_code("Required word deletion flow", "cogs/Sticky.py", 354, 399),
                p(
                    "The required-word delete flow is intentionally conservative. If the bot cannot read history, it avoids deletion to prevent "
                    "false positives. It DMs the thread owner when possible, logs the deletion with author and forum context, unarchives/unlocks if "
                    "needed, and then deletes the thread."
                ),
            ],
        ),
        (
            "Impact And Backup Code Walkthrough",
            [
                p(
                    "The impact system is a reporting pipeline built on top of the bot's existing persistent tables. It does not invent numbers; "
                    "it aggregates workflow records the bot already stores."
                ),
                source_code("Database backup posting", "cogs/Commands.py", 291, 360),
                source_code("Impact metric collection entry", "cogs/Commands.py", 564, 642),
                source_code("Impact report persistence", "cogs/Commands.py", 1106, 1164),
                p(
                    "A backup is a zipped SQLite copy posted to Discord and recorded in database_backups. An impact report is a Markdown/CSV/JSON "
                    "bundle posted to Discord and recorded in impact_snapshots. The two systems serve different purposes: backup is recovery; impact "
                    "is evidence and forecasting."
                ),
                diagram(
                    "Impact Data Path",
                    [
                        "workflow tables",
                        "metric collector",
                        "forecast helper",
                        "Discord attachments",
                        "impact_snapshots row",
                    ],
                    "The report channel gives you durable evidence files; the database row keeps a bot-side history.",
                ),
            ],
        ),
        (
            "Server Icon Rotation Code Walkthrough",
            [
                p(
                    "Server icon rotation is a good example of making a visually simple feature reliable. The feature needs config validation, "
                    "URL cleaning, mode selection, interval enforcement, download checks, state persistence, and commands for manual control."
                ),
                source_code("Icon config normalization", "utils/server_icons.py", 8, 57),
                source_code("Automatic icon rotation loop", "cogs/Background.py", 857, 879),
                source_code("Server icon command surface", "cogs/Commands.py", 1375, 1465),
                p(
                    "The current_index and current_url fields prevent linear rotation from getting stuck and help the bot know what it last tried. "
                    "The interval is normalized to at least five minutes to respect Discord rate limits and avoid accidental rapid icon changes."
                ),
            ],
        ),
        (
            "Engineering Thinking Behind The Bot",
            [
                p(
                    "Avenue Guard's best design choices are about recovering from imperfect reality. Discord is not a database. Messages disappear, "
                    "permissions change, users close DMs, external APIs fail, and hosted storage can be wiped. The bot works because the important "
                    "truth lives in SQLite and visible Discord messages are treated as projections that can be refreshed or rebuilt."
                ),
                table(
                    ["Problem", "Design Response"],
                    [
                        ["A button can be clicked after restart", "Register persistent views with stable custom IDs."],
                        ["Two users can submit at the same time", "Use submit locks and database duplicate constraints."],
                        ["Two reviewers can click the same request", "Use review lock, pending status check, and disabled buttons."],
                        ["A provider can fail or disagree", "Cache normalized validation, warn on uncertainty, and circuit-break repeated failures."],
                        ["A ticket close can fail midway", "Save transcript before deletion and restore status on failure."],
                        ["Render can wipe source storage", "Resolve DB path to persistent disk and post zipped backups."],
                        ["Config can drift from Discord", "Expose dashboard, config check, doctor, repair, and storage commands."],
                    ],
                ),
                p(
                    "The bot is not architected as many isolated mini-bots. It is one coherent system where request data can feed impact reports, "
                    "weekly tracking can feed request reviews, help data can feed support metrics, and background telemetry can feed forecasts."
                ),
                callout(
                    "How we planned it",
                    "The pattern was usually: identify a manual staff pain, decide what state must survive, store that state, expose a Discord UI, "
                    "log the outcome, then add a repair or diagnostic path for the ways Discord can drift.",
                ),
            ],
        ),
        (
            "Debugging Notebook",
            [
                p(
                    "When something breaks, resist the urge to read everything. Use the failure type to choose the shortest path."
                ),
                table(
                    ["Symptom", "First Places To Check"],
                    [
                        ["Bot fails on startup", "Render logs, main.py resolve_db_path, utils/db.py connect/migrate, config.json syntax."],
                        ["Command says no permission", "CommandsCog permission helper, roles.admin_owner_role_ids, MOD_ROLE_ID, reviewer_role_ids."],
                        ["Request button wrong state", "level_request_state row, refresh_or_create_request_button, request channel/message IDs."],
                        ["Request modal rejects valid user", "required_role_ids, request_banned_role_id, _requirements_ok, guild member role cache."],
                        ["ID validation seems wrong", "gd_level_validation_cache, provider settings, external provider output, circuit breaker state."],
                        ["Weekly winner not DMed", "weekly_reward_disabled, weekly_runs, weekly_claims, excluded roles, DM failure log."],
                        ["Ticket status does not update", "tickets.opening_message_id, update_ticket_opening_status, HelpCog on_message."],
                        ["Impact data missing", "daily_stats persistence, impact.allowed_user_ids, impact report channel, database path persistence."],
                    ],
                ),
                h2("The Five Checks"),
                numbered(
                    [
                        "Check config: is the channel/role/table setting pointing at the right thing?",
                        "Check database state: does SQLite say the workflow is open, closed, pending, reviewed, or missing?",
                        "Check Discord object: does the message/channel/thread still exist?",
                        "Check permissions: can the bot see, send, edit, delete, manage roles, or manage channels there?",
                        "Check repair path: does /bot dashboard, /bot doctor, /requests repair, /bot storage, or /bot backup explain the issue?",
                    ]
                ),
                p(
                    "This debugging style matches the code's architecture. Config chooses targets, SQLite stores truth, Discord shows a projection, "
                    "permissions decide whether the projection can be updated, and repair commands rebuild the projection when it drifts."
                ),
            ],
        ),
    ]


CHAPTERS = [
    (
        "Private Manual Notice",
        [
            p(
                "This document is now a private technical manual for Rodrigo. It explains Avenue Guard as a codebase, not as a public "
                "staff guide. It still describes staff-facing workflows because those workflows exist in the code, but the reader is assumed "
                "to be the person trying to understand, maintain, explain, and continue building the bot."
            ),
            callout(
                "Reading promise",
                "The manual uses plain language, diagrams, tables, and real source excerpts. The aim is to make the bot understandable from "
                "top to bottom without flattening the technical details that make it work.",
            ),
            p(
                "If you only need a quick answer, use the inventory and appendix chapters. If you want a deep read, follow the source "
                "walkthrough chapters in order. They move from startup to persistence, then into requests, tracking, help, telemetry, "
                "impact reporting, and debugging."
            ),
        ],
    )
] + CHAPTERS + _private_deep_dive_chapters()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    pgh = cell.paragraphs[0]
    pgh.paragraph_format.space_after = Pt(0)
    pgh.paragraph_format.line_spacing = 1.15
    run = pgh.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.font.color.rgb = color or INK
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(tbl, widths: list[float]) -> None:
    for row in tbl.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def add_para(doc: Document, text: str, style: str = "Normal", bold: bool = False, italic: bool = False) -> None:
    para = doc.add_paragraph(style=style)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.25
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = INK
    run.bold = bold
    run.italic = italic


def add_callout(doc: Document, title: str, text: str) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    set_table_widths(tbl, [6.5])
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = DARK_BLUE
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(text)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10)
    r2.font.color.rgb = INK
    doc.add_paragraph()


def add_manual_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.autofit = False
    tbl.style = "Table Grid"
    usable = 6.5
    if len(headers) == 2:
        widths = [2.0, 4.5]
    elif len(headers) == 3:
        widths = [1.6, 2.45, 2.45]
    else:
        widths = [usable / len(headers)] * len(headers)
    set_table_widths(tbl, widths)
    for idx, header in enumerate(headers):
        cell = tbl.rows[0].cells[idx]
        set_cell_shading(cell, TABLE_FILL)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE)
    for row_values in rows:
        row = tbl.add_row()
        for idx, value in enumerate(row_values):
            set_cell_text(row.cells[idx], value)
    doc.add_paragraph()


def add_manual_diagram(doc: Document, title: str, steps: list[str], note: str = "") -> None:
    add_para(doc, title, bold=True)
    tbl = doc.add_table(rows=1, cols=max(1, len(steps)))
    tbl.autofit = False
    width = 6.5 / max(1, len(steps))
    set_table_widths(tbl, [width] * max(1, len(steps)))
    for idx, step in enumerate(steps):
        cell = tbl.cell(0, idx)
        set_cell_shading(cell, TABLE_FILL)
        set_cell_text(cell, f"{idx + 1}. {step}", bold=True, color=DARK_BLUE)
    if note:
        add_callout(doc, "How to read this", note)
    else:
        doc.add_paragraph()


def add_code_block(doc: Document, title: str, language: str, text: str) -> None:
    add_para(doc, title, bold=True)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    set_table_widths(tbl, [6.5])
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, CODE_FILL)
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(str(text or ""))
    run.font.name = "Courier New"
    run.font.size = Pt(7.3)
    run.font.color.rgb = INK
    if language:
        para2 = cell.add_paragraph()
        para2.paragraph_format.space_after = Pt(0)
        r2 = para2.add_run(f"language: {language}")
        r2.font.name = "Calibri"
        r2.font.size = Pt(7.5)
        r2.font.color.rgb = MUTED
    doc.add_paragraph()


def set_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(60)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Avenue Guard")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.color.rgb = BLUE
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("Private Technical Manual: How The Bot Works Internally")
    run.font.name = "Calibri"
    run.font.size = Pt(15)
    run.font.color.rgb = MUTED

    add_callout(
        doc,
        "Purpose",
        "A private, code-aware handbook for understanding how Avenue Guard works, why each system exists, and how the workflows are implemented.",
    )
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(18)
    r = meta.add_run("Generated from the current Avenue Guard codebase and configuration for Rodrigo.")
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED


def add_toc(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Contents", level=1)
    for idx, (title, _) in enumerate(CHAPTERS, start=1):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(f"{idx}. {title}")
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = INK


def build_docx() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_styles(doc)
    add_cover(doc)
    add_toc(doc)

    for idx, (title, blocks) in enumerate(CHAPTERS, start=1):
        doc.add_page_break()
        doc.add_heading(f"{idx}. {title}", level=1)
        for block in blocks:
            kind = block[0]
            if kind == "p":
                add_para(doc, block[1])
            elif kind == "h2":
                doc.add_heading(block[1], level=2)
            elif kind == "h3":
                doc.add_heading(block[1], level=3)
            elif kind == "bullets":
                for item in block[1]:
                    para = doc.add_paragraph(style="List Bullet")
                    para.paragraph_format.left_indent = Inches(0.375)
                    para.paragraph_format.first_line_indent = Inches(-0.188)
                    para.paragraph_format.space_after = Pt(4)
                    para.paragraph_format.line_spacing = 1.25
                    run = para.add_run(item)
                    run.font.name = "Calibri"
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = INK
            elif kind == "numbered":
                for item in block[1]:
                    para = doc.add_paragraph(style="List Number")
                    para.paragraph_format.left_indent = Inches(0.375)
                    para.paragraph_format.first_line_indent = Inches(-0.188)
                    para.paragraph_format.space_after = Pt(4)
                    para.paragraph_format.line_spacing = 1.25
                    run = para.add_run(item)
                    run.font.name = "Calibri"
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = INK
            elif kind == "callout":
                add_callout(doc, block[1], block[2])
            elif kind == "table":
                add_manual_table(doc, block[1], block[2])
            elif kind == "diagram":
                add_manual_diagram(doc, block[1], block[2], block[3])
            elif kind == "code":
                add_code_block(doc, block[1], block[2], block[3])

    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        header = section.header.paragraphs[0]
        header.text = "Avenue Guard Private Technical Manual"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            run.font.color.rgb = MUTED
        footer = section.footer.paragraphs[0]
        footer.text = "Private code architecture handbook"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            run.font.color.rgb = MUTED

    doc.save(DOCX_PATH)


def build_markdown() -> None:
    lines = [
        "# Avenue Guard",
        "",
        "Private Technical Manual: How The Bot Works Internally",
        "",
        "Generated from the current Avenue Guard codebase and configuration for Rodrigo.",
        "",
        "## Contents",
        "",
    ]
    for idx, (title, _) in enumerate(CHAPTERS, start=1):
        lines.append(f"{idx}. {title}")
    lines.append("")

    for idx, (title, blocks) in enumerate(CHAPTERS, start=1):
        lines.extend([f"## {idx}. {title}", ""])
        for block in blocks:
            kind = block[0]
            if kind == "p":
                lines.extend([block[1], ""])
            elif kind == "h2":
                lines.extend([f"### {block[1]}", ""])
            elif kind == "h3":
                lines.extend([f"#### {block[1]}", ""])
            elif kind == "bullets":
                lines.extend([f"- {item}" for item in block[1]])
                lines.append("")
            elif kind == "numbered":
                lines.extend([f"{n}. {item}" for n, item in enumerate(block[1], start=1)])
                lines.append("")
            elif kind == "callout":
                lines.extend([f"> **{block[1]}:** {block[2]}", ""])
            elif kind == "table":
                headers, rows = block[1], block[2]
                lines.append("| " + " | ".join(headers) + " |")
                lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                for row in rows:
                    safe = [str(cell).replace("\n", "<br>") for cell in row]
                    lines.append("| " + " | ".join(safe) + " |")
                lines.append("")
            elif kind == "diagram":
                title, steps, note = block[1], block[2], block[3]
                lines.extend([f"#### {title}", "", "```mermaid", "flowchart LR"])
                for step_idx, step in enumerate(steps, start=1):
                    safe = str(step).replace('"', "'")
                    lines.append(f'  S{step_idx}["{safe}"]')
                for step_idx in range(1, len(steps)):
                    lines.append(f"  S{step_idx} --> S{step_idx + 1}")
                lines.append("```")
                if note:
                    lines.extend(["", f"> {note}"])
                lines.append("")
            elif kind == "code":
                title, language, text = block[1], block[2], block[3]
                lines.extend([f"#### {title}", "", f"```{language}"])
                lines.extend(str(text or "").splitlines())
                lines.extend(["```", ""])

    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def _pdf_styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ManualTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=28,
            leading=34,
            textColor=colors.HexColor("#1D6B73"),
            alignment=TA_CENTER,
            spaceAfter=12,
        ),
        "subtitle": ParagraphStyle(
            "ManualSubtitle",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#666666"),
            alignment=TA_CENTER,
            spaceAfter=28,
        ),
        "h1": ParagraphStyle(
            "ManualH1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#1D6B73"),
            spaceBefore=8,
            spaceAfter=10,
        ),
        "h2": ParagraphStyle(
            "ManualH2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=17,
            textColor=colors.HexColor("#1D6B73"),
            spaceBefore=10,
            spaceAfter=7,
        ),
        "h3": ParagraphStyle(
            "ManualH3",
            parent=base["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=11.5,
            leading=15,
            textColor=colors.HexColor("#173F46"),
            spaceBefore=8,
            spaceAfter=5,
        ),
        "body": ParagraphStyle(
            "ManualBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10.2,
            leading=13.2,
            textColor=colors.HexColor("#222222"),
            spaceAfter=7,
        ),
        "small": ParagraphStyle(
            "ManualSmall",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8.7,
            leading=11.2,
            textColor=colors.HexColor("#222222"),
        ),
        "callout_title": ParagraphStyle(
            "ManualCalloutTitle",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=9.8,
            leading=12.2,
            textColor=colors.HexColor("#173F46"),
            spaceAfter=2,
        ),
        "callout_body": ParagraphStyle(
            "ManualCalloutBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.2,
            leading=11.8,
            textColor=colors.HexColor("#222222"),
        ),
        "code_title": ParagraphStyle(
            "ManualCodeTitle",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=8.6,
            leading=10.6,
            textColor=colors.HexColor("#173F46"),
            spaceAfter=3,
        ),
        "code": ParagraphStyle(
            "ManualCode",
            parent=base["BodyText"],
            fontName="Courier",
            fontSize=6.2,
            leading=7.4,
            textColor=colors.HexColor("#222222"),
        ),
    }


def _pdf_safe(text: str) -> str:
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _pdf_header_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#666666"))
    canvas.drawString(doc.leftMargin, doc.height + doc.topMargin + 0.18 * inch, "Avenue Guard Private Technical Manual")
    canvas.drawRightString(doc.leftMargin + doc.width, 0.55 * inch, f"Page {doc.page}")
    canvas.restoreState()


def _pdf_callout(styles, title: str, text: str):
    data = [
        [Paragraph(_pdf_safe(title), styles["callout_title"])],
        [Paragraph(_pdf_safe(text), styles["callout_body"])],
    ]
    tbl = Table(data, colWidths=[6.25 * inch])
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFF7E6")),
                ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor("#E6D6B8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return tbl


def _pdf_table(styles, headers: list[str], rows: list[list[str]]):
    data = [[Paragraph(_pdf_safe(h), styles["small"]) for h in headers]]
    for row in rows:
        data.append([Paragraph(_pdf_safe(cell), styles["small"]) for cell in row])
    col_count = max(len(headers), 1)
    if col_count == 2:
        widths = [1.75 * inch, 4.5 * inch]
    elif col_count == 3:
        widths = [1.45 * inch, 2.4 * inch, 2.4 * inch]
    else:
        widths = [(6.25 / col_count) * inch for _ in range(col_count)]
    tbl = Table(data, colWidths=widths, repeatRows=1)
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCEDEA")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#173F46")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B8D0CC")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return tbl


def _pdf_diagram(styles, title: str, steps: list[str], note: str = ""):
    story = [Paragraph(_pdf_safe(title), styles["h3"])]
    if not steps:
        return story
    cell_width = (6.25 / max(1, len(steps))) * inch
    data = [[Paragraph(_pdf_safe(f"{idx + 1}. {step}"), styles["small"]) for idx, step in enumerate(steps)]]
    tbl = Table(data, colWidths=[cell_width] * len(steps))
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#DCEDEA")),
                ("BOX", (0, 0), (-1, -1), 0.55, colors.HexColor("#82A9A2")),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B8D0CC")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.append(tbl)
    if note:
        story.append(Spacer(1, 0.06 * inch))
        story.append(_pdf_callout(styles, "How to read this", note))
    story.append(Spacer(1, 0.12 * inch))
    return story


def _pdf_code_block(styles, title: str, language: str, text: str):
    raw_lines = str(text or "").splitlines() or [""]
    story = []
    chunk_size = 46
    chunks = [raw_lines[idx: idx + chunk_size] for idx in range(0, len(raw_lines), chunk_size)]
    for chunk_idx, chunk in enumerate(chunks, start=1):
        chunk_title = title if len(chunks) == 1 else f"{title} - part {chunk_idx}"
        escaped_lines = [_pdf_safe(line) for line in chunk]
        code_text = "<br/>".join(escaped_lines)
        data = [
            [Paragraph(_pdf_safe(chunk_title), styles["code_title"])],
            [Paragraph(code_text, styles["code"])],
        ]
        tbl = Table(data, colWidths=[6.25 * inch])
        tbl.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F6F7F8")),
                    ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor("#D5D8DC")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.extend([tbl, Spacer(1, 0.11 * inch)])
    return story


def build_pdf() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    styles = _pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.8 * inch,
        leftMargin=0.8 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        title="Avenue Guard Private Technical Manual",
        author="Avenue Guard",
    )
    story = [Spacer(1, 1.1 * inch)]
    story.append(Paragraph("Avenue Guard", styles["title"]))
    story.append(Paragraph("Private Technical Manual: How The Bot Works Internally", styles["subtitle"]))
    story.append(
        _pdf_callout(
            styles,
            "Purpose",
            "A private, code-aware handbook for understanding how Avenue Guard works, why each system exists, and how the workflows are implemented.",
        )
    )
    story.append(Spacer(1, 0.25 * inch))
    story.append(Paragraph("Generated from the current Avenue Guard codebase and configuration for Rodrigo.", styles["subtitle"]))
    story.append(PageBreak())

    story.append(Paragraph("Contents", styles["h1"]))
    for idx, (title, _) in enumerate(CHAPTERS, start=1):
        story.append(Paragraph(f"{idx}. {_pdf_safe(title)}", styles["body"]))
    story.append(PageBreak())

    for idx, (title, blocks) in enumerate(CHAPTERS, start=1):
        story.append(Paragraph(f"{idx}. {_pdf_safe(title)}", styles["h1"]))
        for block in blocks:
            kind = block[0]
            if kind == "p":
                story.append(Paragraph(_pdf_safe(block[1]), styles["body"]))
            elif kind == "h2":
                story.append(Paragraph(_pdf_safe(block[1]), styles["h2"]))
            elif kind == "h3":
                story.append(Paragraph(_pdf_safe(block[1]), styles["h3"]))
            elif kind == "bullets":
                for item in block[1]:
                    story.append(Paragraph(f"- {_pdf_safe(item)}", styles["body"]))
            elif kind == "numbered":
                for number, item in enumerate(block[1], start=1):
                    story.append(Paragraph(f"{number}. {_pdf_safe(item)}", styles["body"]))
            elif kind == "callout":
                story.append(_pdf_callout(styles, block[1], block[2]))
                story.append(Spacer(1, 0.08 * inch))
            elif kind == "table":
                story.append(_pdf_table(styles, block[1], block[2]))
                story.append(Spacer(1, 0.12 * inch))
            elif kind == "diagram":
                story.extend(_pdf_diagram(styles, block[1], block[2], block[3]))
            elif kind == "code":
                story.extend(_pdf_code_block(styles, block[1], block[2], block[3]))
        story.append(PageBreak())

    doc.build(story, onFirstPage=_pdf_header_footer, onLaterPages=_pdf_header_footer)


if __name__ == "__main__":
    build_docx()
    build_markdown()
    build_pdf()
    print(DOCX_PATH)
    print(MD_PATH)
    print(PDF_PATH)
